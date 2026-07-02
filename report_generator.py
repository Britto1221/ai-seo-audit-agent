from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from pathlib import Path
from schemas import FinalSEOReport


OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)


def save_docx_report(report: FinalSEOReport, output_path: str = "outputs/seo_report.docx") -> str:
    doc = Document()

    doc.add_heading("AI SEO Audit Report", level=1)

    doc.add_paragraph(f"Website: {report.website_url}")
    doc.add_paragraph(f"Score: {report.site_score}")
    doc.add_paragraph(f"Grade: {report.grade}")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(report.executive_summary)

    doc.add_heading("Top Problems", level=2)
    for problem in report.top_problems:
        doc.add_paragraph(problem, style="List Bullet")

    doc.add_heading("Technical Summary", level=2)
    doc.add_paragraph(report.technical_summary)

    doc.add_heading("Content Summary", level=2)
    doc.add_paragraph(report.content_summary)

    doc.add_heading("Priority Actions", level=2)
    for action in report.priority_actions:
        doc.add_heading(f"{action.priority}: {action.issue}", level=3)
        doc.add_paragraph(f"Why it matters: {action.why_it_matters}")
        doc.add_paragraph(f"Recommended fix: {action.recommended_fix}")
        doc.add_paragraph(f"Expected impact: {action.expected_impact}")

    doc.add_heading("Page Recommendations", level=2)
    for page in report.page_recommendations:
        doc.add_heading(page.page_url, level=3)
        doc.add_paragraph(page.summary)

        if page.suggested_title:
            doc.add_paragraph(f"Suggested title: {page.suggested_title}")

        if page.suggested_meta_description:
            doc.add_paragraph(f"Suggested meta description: {page.suggested_meta_description}")

        if page.suggested_h1:
            doc.add_paragraph(f"Suggested H1: {page.suggested_h1}")

        for suggestion in page.content_suggestions:
            doc.add_paragraph(suggestion, style="List Bullet")

    doc.add_heading("30-Day SEO Plan", level=2)
    for item in report.thirty_day_plan:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output_path)
    return output_path


def save_pdf_report(report: FinalSEOReport, output_path: str = "outputs/seo_report.pdf") -> str:
    pdf = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    y = height - 50

    def write_line(text: str, size: int = 11, gap: int = 18):
        nonlocal y

        if y < 60:
            pdf.showPage()
            y = height - 50

        pdf.setFont("Helvetica", size)
        pdf.drawString(50, y, str(text)[:100])
        y -= gap

    write_line("AI SEO Audit Report", 18, 30)
    write_line(f"Website: {report.website_url}")
    write_line(f"Score: {report.site_score}")
    write_line(f"Grade: {report.grade}", gap=30)

    write_line("Executive Summary", 14, 24)
    for line in _wrap_text(report.executive_summary):
        write_line(line)

    write_line("Top Problems", 14, 24)
    for problem in report.top_problems:
        for line in _wrap_text(f"- {problem}"):
            write_line(line)

    write_line("Priority Actions", 14, 24)
    for action in report.priority_actions:
        write_line(f"{action.priority}: {action.issue}", 12, 20)
        for line in _wrap_text(f"Fix: {action.recommended_fix}"):
            write_line(line)

    write_line("30-Day SEO Plan", 14, 24)
    for item in report.thirty_day_plan:
        for line in _wrap_text(f"- {item}"):
            write_line(line)

    pdf.save()
    return output_path


def _wrap_text(text: str, max_length: int = 95) -> list[str]:
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        if len(current) + len(word) + 1 <= max_length:
            current += " " + word if current else word
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines